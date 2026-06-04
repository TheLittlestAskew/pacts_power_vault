# Build the Davies Shelters two-page landscape portfolio excerpt (.docx)
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- Brand colors ----
OFFWHITE  = 'FFFDF7'
NAVY      = '112B74'
SLATE     = '2B3B54'
GRAY      = '4B5563'
BLUE      = '1D4ED8'
GREEN     = '16A34A'
BLUETINT  = 'D3E2FF'
GREENTINT = 'B7E4C7'
LIGHTGRAY = 'F1F5F9'

MONT   = 'Montserrat'
BARLOW = 'Barlow'

doc = Document()

# ---- Helpers ----
def style_run(r, font=BARLOW, size=11, color=SLATE, bold=False, italic=False, tracking=None):
    r.font.name = font
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    r.font.bold = bold
    r.font.italic = italic
    rPr = r._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(attr), font)
    if tracking:
        sp = OxmlElement('w:spacing')
        sp.set(qn('w:val'), str(tracking))
        rPr.append(sp)

def para(container, text='', font=BARLOW, size=11, color=SLATE, bold=False,
         italic=False, tracking=None, before=0, after=6, line=1.12,
         align=WD_ALIGN_PARAGRAPH.LEFT, right_indent=None, reuse_first=False):
    if reuse_first and container.paragraphs and not container.paragraphs[0].runs \
            and container.paragraphs[0].text == '':
        p = container.paragraphs[0]
    else:
        p = container.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    p.alignment = align
    if right_indent is not None:
        pf.right_indent = Inches(right_indent)
    if text:
        r = p.add_run(text)
        style_run(r, font=font, size=size, color=color, bold=bold,
                  italic=italic, tracking=tracking)
    return p

def bottom_border(p, color=BLUE, sz=6, space=4):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), str(space))
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def bullet(container, text, size=10.5, color=SLATE, bold=False, after=3, indent=0.18):
    p = container.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(indent)
    pf.first_line_indent = Inches(-indent)
    pf.space_before = Pt(0)
    pf.space_after = Pt(after)
    pf.line_spacing = 1.08
    pf.tab_stops.add_tab_stop(Inches(indent))
    r = p.add_run('•\t' + text)
    style_run(r, font=BARLOW, size=size, color=color, bold=bold)
    return p

def heading(cell_or_doc, text, size=12, before=0, after=7, ruled=True, reuse_first=False):
    p = para(cell_or_doc, text, font=MONT, size=size, color=NAVY, bold=True,
             tracking=8, before=before, after=after, reuse_first=reuse_first)
    if ruled:
        bottom_border(p, color=BLUE, sz=6, space=3)
    return p

def set_cell_bg(cell, color):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shd)

def set_cell_width(cell, inches):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(int(inches * 1440)))
    tcW.set(qn('w:type'), 'dxa')

def fixed_layout(table, widths):
    table.autofit = False
    tblPr = table._tbl.tblPr
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    tblPr.append(layout)
    # strip all borders
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement('w:' + edge)
        el.set(qn('w:val'), 'none')
        borders.append(el)
    tblPr.append(borders)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i])

def set_cell_margins(table, top=180, left=220, bottom=180, right=220):
    tblPr = table._tbl.tblPr
    mar = OxmlElement('w:tblCellMar')
    for tag, val in (('top', top), ('left', left), ('bottom', bottom), ('right', right)):
        el = OxmlElement('w:' + tag)
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        mar.append(el)
    tblPr.append(mar)

def shrink_p(p):
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    pPr = p._p.get_or_add_pPr()
    rPr = OxmlElement('w:rPr')
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '4')  # 2pt — invisible spacer
    rPr.append(sz)
    pPr.append(rPr)

def setup_section(sec):
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Inches(11)
    sec.page_height = Inches(8.5)
    sec.left_margin = Inches(0.65)
    sec.right_margin = Inches(0.65)
    sec.top_margin = Inches(0.55)
    sec.bottom_margin = Inches(0.6)
    sec.footer_distance = Inches(0.32)

def set_footer(sec, text):
    sec.footer.is_linked_to_previous = False
    p = sec.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in list(p.runs):
        r.text = ''
    r = p.add_run(text)
    style_run(r, font=MONT, size=8, color=GRAY, bold=False, tracking=24)

def header_block(eyebrow, title, subtitle, title_size=25):
    para(doc, eyebrow, font=MONT, size=9, color=GRAY, bold=True,
         tracking=36, after=8, reuse_first=True)
    para(doc, title, font=MONT, size=title_size, color=NAVY, bold=True,
         after=4, line=1.05)
    p = para(doc, subtitle, font=BARLOW, size=13, color=GRAY, after=14, line=1.1)
    bottom_border(p, color=BLUE, sz=8, space=8)

# ---- Page background (soft off-white) ----
bg = OxmlElement('w:background')
bg.set(qn('w:color'), OFFWHITE)
doc.element.insert(0, bg)
doc.settings.element.insert(0, OxmlElement('w:displayBackgroundShape'))

# =====================================================================
# PAGE 1
# =====================================================================
sec1 = doc.sections[0]
setup_section(sec1)
set_footer(sec1, 'TECHNOLOGY-ENABLED SHELTER OPERATIONS — OPERATIONS MANUAL EXCERPT')

header_block(
    'NONPROFIT OPERATIONS · SYSTEMS DOCUMENTATION SAMPLE',
    'Technology-Enabled Shelter Operations',
    'Documenting the workflows staff need to capture clean, useful program data',
)

para(doc,
     'As part of the Davies Shelters Operations Manual, I documented the EcoSystem App '
     'procedures that connected shelter operations, case management, shift documentation, '
     'resource coordination, and reporting. The goal was to make daily staff documentation '
     'consistent enough to support service coordination, grant reporting, program '
     'evaluation, and leadership decision-making.',
     font=BARLOW, size=11.5, color=SLATE, after=16, line=1.2, right_indent=1.6)

# Three panels: callout (green tint), data accuracy (blue tint), reporting (light gray)
t1 = doc.add_table(rows=1, cols=5)
fixed_layout(t1, [3.9, 0.2, 2.7, 0.2, 2.7])
set_cell_margins(t1, top=200, left=240, bottom=200, right=240)

c_callout, c_gap1, c_data, c_gap2, c_report = t1.rows[0].cells

# Panel 1 — What This Demonstrates (main callout, green tint)
set_cell_bg(c_callout, GREENTINT)
heading(c_callout, 'What This Demonstrates', size=13, after=9, reuse_first=True)
para(c_callout,
     'This documentation turned a custom internal app into a usable staff workflow. '
     'It connected daily data entry to case management, resource coordination, grant '
     'reporting, board reporting, program evaluation, and funding justification. The '
     'result was not just a system staff could use, but a repeatable documentation '
     'standard that made program data more reliable and funder-ready.',
     font=BARLOW, size=11, color=SLATE, after=0, line=1.25)

# Panel 2 — Data Accuracy Importance (blue tint)
set_cell_bg(c_data, BLUETINT)
heading(c_data, 'Data Accuracy Importance', size=12, after=8, reuse_first=True)
para(c_data, 'The information entered in EcoSystem is used for:',
     font=BARLOW, size=10.5, color=SLATE, after=6, line=1.15)
for item in ('Guest services coordination', 'Grant reporting and compliance',
             'Funding applications', 'Resource allocation', 'Performance metrics'):
    bullet(c_data, item)
para(c_data,
     'Accurate, complete data supports stronger reporting, better decisions, '
     'and better services for guests.',
     font=BARLOW, size=10.5, color=GREEN, bold=True, before=6, after=0, line=1.2)

# Panel 3 — Reporting Functions (light gray)
set_cell_bg(c_report, LIGHTGRAY)
heading(c_report, 'Reporting Functions', size=12, after=8, reuse_first=True)
para(c_report, 'EcoSystem includes reporting capabilities for:',
     font=BARLOW, size=10.5, color=SLATE, after=6, line=1.15)
for item in ('Guest demographics', 'Length of stay statistics', 'Case meeting frequency',
             'Resource utilization', 'Discharge outcomes', 'Phase advancement rates'):
    bullet(c_report, item)
para(c_report, 'Reports are used for:', font=BARLOW, size=10.5, color=SLATE,
     bold=True, before=6, after=6, line=1.15)
for item in ('Grant applications and reporting', 'Board presentations',
             'Program evaluation', 'Funding justification'):
    bullet(c_report, item, after=3 if item != 'Funding justification' else 0)

shrink_p(doc.add_paragraph())

# =====================================================================
# PAGE 2
# =====================================================================
sec2 = doc.add_section(WD_SECTION.NEW_PAGE)
setup_section(sec2)
set_footer(sec2, 'ECOSYSTEM APP WORKFLOW STANDARDS — OPERATIONS MANUAL EXCERPT')

header_block(
    'SELECTED SOP EXCERPT · ECOSYSTEM APP DOCUMENTATION',
    'EcoSystem App Workflow Standards',
    'Selected procedures for case meetings, resource coordination, and shift documentation',
    title_size=23,
)

t2 = doc.add_table(rows=1, cols=5)
fixed_layout(t2, [3.1, 0.2, 3.1, 0.2, 3.1])
set_cell_margins(t2, top=180, left=220, bottom=180, right=220)
col1, gap1, col2, gap2, col3 = t2.rows[0].cells

# ---- Column 1 ----
set_cell_bg(col1, LIGHTGRAY)
heading(col1, 'EcoSystem App Structure', size=11.5, after=7, reuse_first=True)
para(col1,
     'When staff open the app, they see the shelter calendar. The calendar displays '
     'scheduled case meetings and can expand to include scheduled workers, staff '
     'meetings, and volunteer events.',
     font=BARLOW, size=10.5, color=SLATE, after=4, line=1.18)

heading(col1, 'Navigation', size=11.5, before=10, after=7)
para(col1, 'The navigation pane displays the areas each user has access to, including:',
     font=BARLOW, size=10.5, color=SLATE, after=5, line=1.15)
for item in ('Guests', 'Case Meetings', 'Budgets', 'Resources', 'Night Watch', 'Reports'):
    bullet(col1, item)

heading(col1, 'Views vs. Forms', size=11.5, before=10, after=7)
para(col1,
     'Views display lists of records with basic information. Active views show current '
     'records; inactive views show historical records and discharged guests.',
     font=BARLOW, size=10.5, color=SLATE, after=6, line=1.18)
para(col1,
     'Forms are used for detailed information input and editing. Staff click '
     '“New” to create a record or select an existing record to open the form.',
     font=BARLOW, size=10.5, color=SLATE, after=0, line=1.18)

# ---- Column 2 ----
set_cell_bg(col2, LIGHTGRAY)
heading(col2, 'Case Meeting Procedures', size=11.5, after=7, reuse_first=True)
para(col2, 'Case meeting records document:',
     font=BARLOW, size=10.5, color=SLATE, after=5, line=1.15)
for item in ('Guest wins or accomplishments since the last meeting', 'Challenges discussed',
             'Goals set', 'Resources provided or coordinated', 'Action items',
             'Phase advancement discussion, if applicable'):
    bullet(col2, item)
para(col2,
     'After the meeting is documented, staff save the record and schedule the next '
     'meeting on the calendar.',
     font=BARLOW, size=10.5, color=SLATE, before=3, after=4, line=1.18)

heading(col2, 'Tracking Guest Wins', size=11.5, before=10, after=7)
para(col2,
     'Guest wins are important for celebrating progress, grant reporting, building '
     'guest confidence, and demonstrating program effectiveness.',
     font=BARLOW, size=10.5, color=SLATE, after=6, line=1.18)
para(col2, 'Document all wins, even small ones:',
     font=BARLOW, size=10.5, color=SLATE, bold=True, after=5, line=1.15)
for item in ('Obtained ID', 'Attended meeting', 'Completed chore without reminder',
             'Made progress on goal', 'Practiced new skill'):
    bullet(col2, item, after=3 if item != 'Practiced new skill' else 0)

# ---- Column 3 ----
set_cell_bg(col3, LIGHTGRAY)
heading(col3, 'Resource Assignment Procedures', size=11.5, after=7, reuse_first=True)
para(col3, 'When connecting a guest with resources, staff document:',
     font=BARLOW, size=10.5, color=SLATE, after=5, line=1.15)
for item in ('Resource type', 'Agency or provider', 'Contact information',
             'Referral date', 'Follow-up requirements', 'Outcome'):
    bullet(col3, item)
para(col3, 'Resource assignments are tracked in case notes.',
     font=BARLOW, size=10.5, color=SLATE, before=3, after=4, line=1.18)

heading(col3, 'Night Watch Documentation', size=11.5, before=10, after=7)
para(col3,
     'Night Watch replaced the traditional paper Daily Log system. Evening, overnight, '
     'and weekend staff use Night Watch to document significant events during each shift.',
     font=BARLOW, size=10.5, color=SLATE, after=6, line=1.18)
para(col3, 'Event types include:', font=BARLOW, size=10.5, color=SLATE, after=5, line=1.15)

nw = col3.add_table(rows=6, cols=2)
fixed_layout(nw, [1.32, 1.32])
left_items = ('Daily Log', 'Phone Calls', 'Visitors', 'Guest Interactions', 'Meals', 'Donations')
right_items = ('Medication Log', 'Food / Item Requests', 'Safety / Emergencies',
               'Intakes / Discharges', 'Staff to Admin', '')
for i in range(6):
    for j, item in ((0, left_items[i]), (1, right_items[i])):
        cell = nw.rows[i].cells[j]
        if item:
            p = bullet(cell, item, size=10, after=2, indent=0.15)
            # bullet() adds a new paragraph; remove the default empty first one
            first = cell.paragraphs[0]
            if not first.runs and first.text == '':
                first._p.getparent().remove(first._p)
        else:
            shrink_p(cell.paragraphs[0])
shrink_p(col3.paragraphs[-1])

shrink_p(doc.add_paragraph())

# ---- Full-width Documentation Standards panel (blue tint) ----
t3 = doc.add_table(rows=1, cols=1)
fixed_layout(t3, [9.7])
set_cell_margins(t3, top=180, left=240, bottom=180, right=240)
c_std = t3.rows[0].cells[0]
set_cell_bg(c_std, BLUETINT)
heading(c_std, 'Documentation Standards', size=11.5, after=7, reuse_first=True)
para(c_std,
     'Staff are instructed to document events as they happen, be specific and objective, '
     'note times when relevant, document both positive and concerning events, flag items '
     'needing follow-up, and review previous entries at the start of shift.',
     font=BARLOW, size=10.5, color=SLATE, after=0, line=1.2)

shrink_p(doc.add_paragraph())

out = r'C:\Users\theli\pacts_power_vault\DND_Sources\Davies Shelters - Operations Manual Portfolio Excerpt.docx'
doc.save(out)
print('Saved:', out)
